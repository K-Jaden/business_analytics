"""Word 파일 LEGO 아이템 교체 + 빨간 사유 추가"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from copy import deepcopy

DOCX_PATH = 'exec_plan_updated.docx'
doc = Document(DOCX_PATH)

RED = RGBColor(0xFF, 0x00, 0x00)

REASON_TEXT = (
    "※ [교체 사유] 기존 5종(AT-AT 75313, Taj Mahal 10256, Home Alone 21330, "
    "Stranger Things 75810, Haunted House 10273)은 Google Trends 검색량이 "
    "48개월 중 대부분 0으로, CH1(검색 관심도) 신호를 확보할 수 없음. "
    "리셀 시장 활성도와 검색량이 높은 아이템으로 전면 교체함: "
    "Millennium Falcon 75192 / Eiffel Tower 10307 / Back to the Future 10300 / "
    "Hogwarts Castle 71043 / Ferrari Daytona SP3 42143."
)

def add_red_paragraph_after(ref_para, text):
    """ref_para 바로 뒤에 빨간 단락 삽입"""
    new_p = OxmlElement('w:p')
    ref_para._element.addnext(new_p)
    # 문서에서 새 단락 찾아서 run 추가
    for para in doc.paragraphs:
        if para._element is new_p:
            run = para.add_run(text)
            run.font.color.rgb = RED
            run.font.bold = True
            run.font.size = Pt(10)
            return para
    return None

def add_red_run_to_para(para, text):
    run = para.add_run(text)
    run.font.color.rgb = RED
    run.font.bold = True
    return run

# ── 1. 단락 수정 ─────────────────────────────────────────────────────────────
# Para[11]: Stranger Things(75810) → Hogwarts Castle (71043)
# Para[12]: Haunted House(10273) → Ferrari Daytona SP3 (42143)

para11 = doc.paragraphs[11]
para12 = doc.paragraphs[12]

print(f'Before Para[11]: {para11.text}')
print(f'Before Para[12]: {para12.text}')

for run in para11.runs:
    run.text = (run.text
        .replace('Stranger Things(75810)', 'Hogwarts Castle (71043)')
        .replace('Stranger Things 75810', 'Hogwarts Castle 71043')
        .replace('75810', '71043'))

for run in para12.runs:
    run.text = (run.text
        .replace('Haunted House(10273)', 'Ferrari Daytona SP3 (42143)')
        .replace('Haunted House(10273)', 'Ferrari Daytona SP3 (42143)')
        .replace('Haunted House', 'Ferrari Daytona SP3')
        .replace('10273', '42143'))

# 단락 텍스트가 runs에 분산된 경우 직접 처리
if 'Stranger Things' in para11.text or '75810' in para11.text:
    for run in para11.runs:
        run.text = run.text.replace('Stranger Things', 'Hogwarts Castle').replace('75810', '71043')

if 'Haunted House' in para12.text or '10273' in para12.text:
    for run in para12.runs:
        run.text = run.text.replace('Haunted House', 'Ferrari Daytona SP3').replace('10273', '42143')

print(f'After  Para[11]: {para11.text}')
print(f'After  Para[12]: {para12.text}')

# Para[12] 뒤에 빨간 사유 단락 삽입
inserted = add_red_paragraph_after(para12, REASON_TEXT)
print(f'빨간 사유 단락 삽입: {"성공" if inserted else "실패"}')

# ── 2. Table[38] 수정 ────────────────────────────────────────────────────────
table = doc.tables[38]

# Row[8]: lego_hogwarts (75810) → lego_hogwarts (71043)
# Row[9]: lego_ferrari (10273) → lego_ferrari (42143)
for row_idx, old_num, new_num, new_name in [
    (8, '75810', '71043', 'lego_hogwarts (71043)\n[교체: Hogwarts Castle 71043]'),
    (9, '10273', '42143', 'lego_ferrari (42143)\n[교체: Ferrari Daytona SP3 42143]'),
]:
    cell = table.rows[row_idx].cells[0]
    print(f'Before Table[38][{row_idx}][0]: {cell.text}')
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = run.text.replace(old_num, new_num)
    # 세트 번호 직접 치환 (run 분산 대비)
    if old_num in cell.text:
        for para in cell.paragraphs:
            full = para.text
            if old_num in full:
                for run in para.runs:
                    run.text = run.text.replace(old_num, new_num)
    print(f'After  Table[38][{row_idx}][0]: {cell.text}')

# Row[8] Col[1] (설명)에도 교체 사유 빨간 텍스트 추가
for row_idx, reason in [
    (8, '※ 교체: Hogwarts Castle 71043 (기존 Stranger Things 75810 — 검색량 0)'),
    (9, '※ 교체: Ferrari Daytona SP3 42143 (기존 Haunted House 10273 — 검색량 0)'),
]:
    cell_reason = table.rows[row_idx].cells[1]
    # 기존 텍스트 뒤에 빨간 텍스트 추가
    para = cell_reason.paragraphs[-1]
    run = para.add_run('\n' + reason)
    run.font.color.rgb = RED
    run.font.bold = True

# ── 3. 새 교체 아이템에 대한 연속성 확인 행 추가 ──────────────────────────
# Eiffel Tower(10307, 2022-10 출시), Back to the Future(10300, 2022-04 출시)
# 는 2022년 중간 출시라 추가 확인 필요
new_check_items = [
    ('lego_eiffel (10307)',   'Eiffel Tower 10307 — 2022-10 출시, 2022-01~09 결측 예상', '교체 아이템, 출시일 이후 데이터만 사용'),
    ('lego_delorean (10300)', 'Back to the Future 10300 — 2022-04 출시, 2022-01~03 결측 예상', '교체 아이템, 출시일 이후 데이터만 사용'),
]

for item_id, desc, action in new_check_items:
    row = table.add_row()
    cells = row.cells
    # Col 0: item_id (빨간색)
    run0 = cells[0].paragraphs[0].add_run(item_id)
    run0.font.color.rgb = RED
    run0.font.bold = True
    # Col 1: desc (빨간색)
    run1 = cells[1].paragraphs[0].add_run(desc)
    run1.font.color.rgb = RED
    run1.font.bold = True
    # Col 2: action (빨간색)
    run2 = cells[2].paragraphs[0].add_run(action)
    run2.font.color.rgb = RED
    run2.font.bold = True
    print(f'새 행 추가: {item_id}')

doc.save(DOCX_PATH)
print(f'\n저장 완료: {DOCX_PATH}')
