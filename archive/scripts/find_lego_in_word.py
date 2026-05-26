"""Word 파일에서 LEGO 관련 내용 찾기"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement

doc = Document('exec_plan_updated.docx')

keywords = ['lego_','75313','10256','21330','75810','10273',
            '75192','10307','10300','71043','42143',
            'AT-AT','Taj','Home Alone','Stranger','Haunted',
            'Millennium','Eiffel','DeLorean','Hogwarts','Ferrari Daytona',
            'falcon','eiffel','delorean','hogwarts','ferrari']

print('=== 단락 검색 ===')
found_paras = []
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if any(k.lower() in t.lower() for k in keywords):
        print(f'Para[{i}]: {t[:150]}')
        found_paras.append(i)

print(f'\n단락 발견: {len(found_paras)}개')

print('\n=== 테이블 검색 ===')
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            t = cell.text
            if any(k.lower() in t.lower() for k in keywords):
                print(f'Table[{ti}] Row[{ri}] Col[{ci}]: {t[:150]}')
